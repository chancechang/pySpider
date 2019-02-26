import xlrd
from bs4 import BeautifulSoup
import requests
import time
import datetime
import json
import re
import xlsxwriter
import os
from docx import Document
from docx.shared import Inches


cook='UM_distinctid=16822772df50-0cca678828aea6-b353461-100200-16822772e8fd5; CNZZDATA5626149=cnzz_eid%3D569227167-1546763274-%26ntime%3D1546779612'
header={
    'Cookie':cook,
    'Host':'m.my478.com',
    'Referer':'http://m.my478.com/jishu/list/',
    # 'Upgrade-Insecure-Requests':'1',
    'User-Agent':'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/64.0.3282.186 Safari/537.36',
}
#种植技术
def get_tv_link_one():
    f=open('tv_link.txt','a+',encoding='utf-8')
    for i in range(1,798):
        print(str(i)+'页')
        # http://m.1988.tv/GetPage.aspx?type=bingchonghai&page=2
        url='http://m.1988.tv/GetPage.aspx?type=bingchonghai&page='+str(i)
        req=requests.get(url)
        req.encoding='gb2312'
        bsObj=BeautifulSoup(req.text,'html.parser')
        jshu=bsObj.find_all('li')
        print(len(jshu))
        for li in jshu:
            a=li.find('a')
           
            f.write(str([a.attrs['href'],a.attrs['title']])+'\n')
            i=i+1
        if len(jshu)==0:
            break
    f.close()
# get_tv_link_one()

def get_Obj(url):
    for i in range(5):
        try:
            req=requests.get(url)
            bsObj=BeautifulSoup(req.text,'html.parser')
            div=bsObj.find('div',class_='list_content')
            li_list=div.find_all('a',recursive=False)
            return li_list
        except:
            pass
        if i == 4:
            print(url+'失效')
    return None

def get_url(type,f):
    for i in range(1,1000):
        print(str(i)+'页')
        # http://m.1988.tv/information/insects/list.html?code=6325&page=2
        url='http://m.1988.tv'+str(type)+'&page='+str(i)
        a_list=get_Obj(url)
        print(len(a_list))
        for a in a_list:            
            href=a.attrs['href']
            title=a.find('h3').get_text().strip()
            row=[href,title]
            f.write(str(row)+'\n')
        if len(a_list)<10:
            break



def veg_get_all_link():   
    f=open('tv_link_two.txt','w+',encoding='utf-8')
    for line in open('tv_link.txt','r',encoding='utf-8'):
        line=eval(line)
        get_url(line[0],f)
    f.close()
# veg_get_all_link()


def d_load(src,imgrootname):
    root='./'+imgrootname
    path='./'+imgrootname+'/'+src.split('/')[-1]
    try:
        if not os.path.exists(path):
            if 'http' in src:
                r = requests.get(src)
            else:
                r = requests.get('http://m.1988.tv'+src)
            r.raise_for_status()
            # 使用with语句可以不用自己手动关闭已经打开的文件流
            with open(path, "wb") as f:  # 开始写文件，wb代表写二进制文件
                f.write(r.content)
            print("爬取完成")
        else:
            print("文件已存在")
    except Exception as e:
        print("爬取失败:"+str(e))
    
    return src.split('/')[-1]
def remove_control_characters(html):
    def str_to_int(s, default, base=10):
        if int(s, base) < 0x10000:
            return unichr(int(s, base))
        return default
        
    html = re.sub(u'&#(\d+);?'  ,lambda c: str_to_int(c.group(1), c.group(0)), html)

    html = re.sub(u"&#[xX]([0-9a-fA-F]+);?", lambda c: str_to_int(c.group(1), c.group(0), base=16), html)
    html = re.sub(u"[\x00-\x08\x0b\x0e-\x1f\x7f]", "", html)
    return html


def get_detail():
    f_error=open('tv_error.txt','a+',encoding='utf-8')
    f_img=open('tv_img.txt','a+',encoding='utf-8')
    # f_final=open('tv_final.txt','a+',encoding='utf-8')
    root='http://m.1988.tv'
    wordrootname='1988_病虫害word'
    htmlrootname='1988_病虫害html'
    imgrootname='1988_病虫害img'
    if not os.path.exists(wordrootname):
        os.mkdir(wordrootname)
    if not os.path.exists(htmlrootname):
        os.mkdir(htmlrootname)
    if not os.path.exists(imgrootname):
        os.mkdir(imgrootname)
    for line in open('tv_final.1.txt','r',encoding='utf-8'):
        line=eval(line)
        fname=line[1].replace('?','').replace('|','').replace('"','').replace('>','').replace('<','').replace('*','').replace('*','').replace('\\','').replace(':','').replace('/','')       
        print(fname)    
        if 'baik' in line[0]:
            path1='./'+wordrootname+'/'+fname+'_'+line[0][7:].replace('.html','')
            path2='./'+htmlrootname+'/'+fname+'_'+line[0][7:].replace('.html','')
        else:
            path1='./'+wordrootname+'/'+fname+'_'+line[0][14:].replace('.html','')            
            path2='./'+htmlrootname+'/'+fname+'_'+line[0][14:].replace('.html','')            
        # print(path)
        # if os.path.exists(path1+'.docx'):
        #     print(path1+'已存在')
        #     continue
        # url=root+line[0] 
        # url='http://www.tv.com/newshow.asp?id=50163'
        # line[1]='白菜要丰产施肥必须有技法'
        # print(url)   
        doc = Document()  
        # for i in range(3):
        #     try:
        #         req=requests.get(url,timeout=15)
        #         req.encoding='utf-8'
        #         bsObj=BeautifulSoup(req.text,'html.parser')
        #         #病虫害
        #         content=bsObj.find('div',class_='content')
        #         img_list=content.find_all('img')
        #         #写入标题
        #         break
        #     except:
        #         pass
        # if content==None  :
        #     print(url+'未取到数据')
        #     f_error.write(str(line)+'\n')
        #     continue
        # if len(img_list)>0:
        #     print(url+'存在图片')
        #     f_img.write(str(line)+'\n')
        #     continue
        line[1]=line[1].replace('\r\n','').strip()
        content=BeautifulSoup(line[2].replace('</p','\n</p'),'html.parser')
        img_list=content.find_all('mip-img')
        # f_final.write(str(line+[str(content)])+'\n')
        # continue        
        

            

        doc.add_heading(line[1],level=0)
        skin_text=str(content)
        for im in img_list:
            src=im.attrs['src']
            print(src)
            newSrc=d_load(src,imgrootname)
            #替换内容
            skin_text=skin_text.replace(src,'../'+imgrootname+'/'+newSrc)
            try:
                doc.add_picture('./'+imgrootname+'/'+newSrc, width=Inches(3)) 
            except:
                pass
            
        #写入文本内容
        # content_word=remove_control_characters(str(content.get_text()))   
        doc.add_paragraph(u'%s'%(content.get_text())) 
        doc.save(path1+'.docx') 
    #写入html    
        
        fp=open(path2+'.html','w',encoding='utf-8')
        fp.write('<html><head><meta http-equiv="Content-Type" content="text/html; charset=utf-8"></head><body><h1>'+str(line[1])+'</h1>'+str(skin_text)+'</body></html>') #写入数据
        fp.close()
        # break
    f_error.close()
    f_img.close()
    # f_final.close()

get_detail()
def asd():
    wordrootname='1988_病虫害word'
    htmlrootname='1988_病虫害html'
    imgrootname='1988_病虫害img'
    rootdir = './1988_病虫害html'
    if not os.path.exists(wordrootname):
        os.mkdir(wordrootname)
    list = os.listdir(rootdir) #列出文件夹下所有的目录与文件
    for i in range(0,len(list)):
        print(i)        
        print(list[i])
        html_path = os.path.join(rootdir,list[i])
        fname=list[i].replace('.html','')
        fname_re=fname.split('_')[0]
        if os.path.isfile(html_path):
            doc = Document()
            doc.add_heading(fname,level=0)
            tex=open(html_path,'r',encoding='utf-8')
            bsObj=BeautifulSoup(tex.read(),'html.parser')
            img_list=bsObj.find_all('img')
            content=str(bsObj).replace('</p','\n</p')
            for  im in img_list:
                try:
                    print(im.attrs['src'][1:])
                    
                    doc.add_picture(im.attrs['src'][1:], ) 
                except:
                    print(os.path.isfile(im.attrs['src'][1:]))
            content_word=remove_control_characters(BeautifulSoup(content,'html.parser').get_text())
            doc.add_paragraph(u'%s'%(content_word.replace(fname_re,''))) 
            doc.save('./'+wordrootname+'/'+fname+'.docx') 
        break       
            
# asd()

