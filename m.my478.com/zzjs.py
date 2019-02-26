
from bs4 import BeautifulSoup
import requests
import time
import datetime
import json
import re
import xlsxwriter
import os
from docx import Document
from docx.shared import Inches


cook='UM_distinctid=16822772df50-0cca678828aea6-b353461-100200-16822772e8fd5; CNZZDATA5626149=cnzz_eid%3D569227167-1546763274-%26ntime%3D1546779612'
header={
    'Cookie':cook,
    'Host':'m.my478.com',
    'Referer':'http://m.my478.com/jishu/list/',
    # 'Upgrade-Insecure-Requests':'1',
    'User-Agent':'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/64.0.3282.186 Safari/537.36',
}
#种植技术
def get_zzjs_link_one():
    url='http://m.my478.com/html/list/'
    f=open('zzjs_line_one.txt','w+',encoding='utf-8')
    req=requests.get(url)
    req.encoding='utf-8'
    bsObj=BeautifulSoup(req.text,'html.parser')
    jshu=bsObj.find('div',class_='jishu').find_all('a')
    print(len(jshu))
    for a in jshu:
        f.write(str([a.attrs['href'],a.get_text()])+'\n')
    f.close()

#种植技术
#问答栏目
#实用技术
#实用技术
#农机常识
def get_zzjs_link_two():
    root='http://m.my478.com'
    f1=open('zzjs_line_two.txt','a+',encoding='utf-8')
    for line in open('zzjs_line_one.txt','r',encoding='utf-8'):
        line=eval(line)
        url=root+line[0]
        url='http://m.my478.com/nongji/list/'
        print(url)
        req=requests.get(url)
        req.encoding='utf-8'
        bsObj=BeautifulSoup(req.text,'html.parser')
        li_list=bsObj.find('div',class_='list-box margin-t5').find_all('li')
        for li in li_list:
            a=li.find('a')
            row=[a.attrs['href'],a.get_text()]
            f1.write(str(row)+'\n')
        break
    f1.close()


def get_bcfz_link():
    f1=open('bcfz_line_two.txt','w+',encoding='utf-8')
    for i in range(1,100):
        print(i)
        url='http://m.my478.com/more/get_bingchong.php?page='+str(i)
        req=requests.get(url,headers=header)
        req.encoding='utf-8'
        bsObj=BeautifulSoup(req.text,'html.parser')
        li_list=bsObj.find_all('li',class_='linked')
        print(len(li_list))
        for li in li_list:
            h3=li.find('h3')
            a=li.find('a').attrs['href']
            row=[a,h3]
            f1.write(str(row)+'\n')
        if len(li_list)==0:
            break
    f1.close()


#病虫防治  http://m.my478.com/more/get_bingchong.php?page=5




def get_link_one():
    f=open('zzjs_line_one.txt','a+',encoding='utf-8')
    url='http://www.vegnet.com.cn/Disease/List_c1025.html'
    req=requests.get(url,headers=header)
    bsObj=BeautifulSoup(req.text,'html.parser')
    a_list=bsObj.find('div',class_='lb2').find_all('a')
    print(len(a_list))
    for a in a_list:
        f.write(str([a.attrs['href'],a.get_text()])+'\n')
    f.close()

def get_Obj(url):
    for i in range(5):
        try:
            req=requests.get(url,headers=header)
            bsObj=BeautifulSoup(req.text,'html.parser')
            div=bsObj.find('ul',class_='bot_list')
            li_list=div.find_all('li')
            return li_list
        except:
            pass
        if i == 4:
            print(url+'失效')
    return None

def get_url(type,f):
    for i in range(1,1000):
        print(str(i)+'页')
        url='http://www.vegnet.com.cn'+str(type)+'_p'+str(i)+'.html'
        p_list=get_Obj(url)
        print(len(p_list))
        for p in p_list:
            a=p.find('a')
            href=a.attrs['href']
            title=a.attrs['title']
            row=[href,title]
            f.write(str(row)+'\n')
        if len(p_list)<12:
            break



def veg_get_all_link():   
    f=open('zzjs_line_two.txt','a+',encoding='utf-8')
    for line in open('zzjs_line_one.txt','r',encoding='utf-8'):
        line=eval(line)

        get_url(line[0].replace('.html',''),f)
    f.close()



def d_load(src,imgrootname):
    root='./'+imgrootname
    path='./'+imgrootname+'/'+src.split('/')[-1]
    try:
        if not os.path.exists(path):
            r = requests.get(src)
            r.raise_for_status()
            # 使用with语句可以不用自己手动关闭已经打开的文件流
            with open(path, "wb") as f:  # 开始写文件，wb代表写二进制文件
                f.write(r.content)
            print("爬取完成")
        else:
            print("文件已存在")
    except Exception as e:
        print("爬取失败:"+str(e))
    
    return src.split('/')[-1]
def remove_control_characters(html):
    def str_to_int(s, default, base=10):
        if int(s, base) < 0x10000:
            return unichr(int(s, base))
        return default
        
    html = re.sub(u'&#(\d+);?'  ,lambda c: str_to_int(c.group(1), c.group(0)), html)

    html = re.sub(u"&#[xX]([0-9a-fA-F]+);?", lambda c: str_to_int(c.group(1), c.group(0), base=16), html)
    html = re.sub(u"[\x00-\x08\x0b\x0e-\x1f\x7f]", "", html)
    return html


def get_detail():
    f_error=open('zzjs_error.txt','w+',encoding='utf-8')
    f_img=open('zzjs_img.txt','w+',encoding='utf-8')
    # f_final=open('zzjs_final1.txt','a+',encoding='utf-8')
    root='http://m.my478.com'
    wordrootname='种植问答网word'
    htmlrootname='种植问答网html'
    imgrootname='种植问答网img'
    if not os.path.exists(wordrootname):
        os.mkdir(wordrootname)
    if not os.path.exists(htmlrootname):
        os.mkdir(htmlrootname)
    if not os.path.exists(imgrootname):
        os.mkdir(imgrootname)
    for line in open('zzjs_final1.txt','r',encoding='utf-8'):
        line=eval(line)
        url=root+line[0] 
        print(url)   
        # content=BeautifulSoup(line[2],'html.parser')
        doc = Document()  
        for i in range(1):
            try:
                # req=requests.get(url,headers=header,timeout=15)
                # req.encoding='utf-8'
                bsObj=BeautifulSoup(line[2],'lxml')
                #病虫害
                content=bsObj.find('div',class_='txt')
                skin_green=bsObj.find('div',class_='img center')
                skin_text=str(skin_green)
                if skin_green==None:
                    img_list=[]
                else:
                    img_list=skin_green.find_all('img')
                #写入标题
                break
            except:
                pass
        if content==None  :
            print(url+'未取到数据')
            f_error.write(str(line)+'\n')
            continue
        # f_final.write(str(line+[str(bsObj)])+'\n')
        # continue
        content_word=str(content).replace('</p','\n</p')
        fname=line[1].strip().replace('?','').replace('|','').replace('"','').replace('>','').replace('<','').replace('*','').replace('*','').replace('\\','').replace(':','').replace('/','')       
        path1='./'+wordrootname+'/'+fname+'_'+line[0].split('/')[-1][:-4]
        path2='./'+htmlrootname+'/'+fname+'_'+line[0].split('/')[-1][:-4]
        # print(path1)
        # break
        
        print(fname)    
        doc.add_heading(line[1],level=0)
        #写入图片
        
        for im in img_list:
            src=im.attrs['src']
            print(src)
            newSrc=d_load(src,imgrootname)
            #替换内容
            skin_text=skin_text.replace(src,'../'+imgrootname+'/'+newSrc)
            try:
                doc.add_picture('./'+imgrootname+'/'+newSrc, width=Inches(3)) 
            except:
                pass
            
        #写入文本内容
        content_word=remove_control_characters(BeautifulSoup(content_word,'html.parser').get_text())   
        doc.add_paragraph(u'%s'%( content_word)) 
        doc.save(path1+'.docx') 
    #写入html    
        
        fp=open(path2+'.html','w',encoding='utf-8')
        if skin_green==None:
            fp.write('<html><head><meta http-equiv="Content-Type" content="text/html; charset=utf-8"></head><body><h1>'+str(line[1])+'</h1>'+str(content)+'</body></html>') #写入数据
        else:
            fp.write('<html><head><meta http-equiv="Content-Type" content="text/html; charset=utf-8"></head><body><h1>'+str(line[1])+'</h1>'+str(skin_text)+str(content)+'</body></html>') #写入数据
            
        fp.close()
        # break
    f_error.close()
    f_img.close()
    # f_final.close()

get_detail()


# url='http://m.my478.com/html/201808/118096.htm'
# req=requests.get(url)
# req.encoding='utf-8'
# bsObj=BeautifulSoup(req.text,'lxml')
# #病虫害
# content=bsObj.find('div',class_='txt')
# print(content)

def asd():
    wordrootname='种植问答网word'
    htmlrootname='种植问答网html'
    imgrootname='种植问答网img'
    rootdir = './种植问答网html'
    if not os.path.exists(wordrootname):
        os.mkdir(wordrootname)
    list = os.listdir(rootdir) #列出文件夹下所有的目录与文件
    for i in range(111,len(list)):
        print(i)
        
        print(list[i])
        html_path = os.path.join(rootdir,list[i])
        fname=list[i].replace('.html','')
        if os.path.isfile(html_path):
            doc = Document()
            doc.add_heading(fname,level=0)
            tex=open(html_path,'r',encoding='utf-8')
            bsObj=BeautifulSoup(tex.read(),'html.parser')
            img_list=bsObj.find_all('img')
            content=str(bsObj).replace('</p','\n\r</p')
            for  im in img_list:
                try:
                    print(im.attrs['src'][1:])
                    
                    doc.add_picture(im.attrs['src'][1:], ) 
                except:
                    print(os.path.isfile(im.attrs['src'][1:]))
            content_word=remove_control_characters(BeautifulSoup(content,'html.parser').get_text())
            doc.add_paragraph(u'%s'%(content_word.replace(fname,''))) 
            doc.save('./'+wordrootname+'/'+fname+'.docx') 
                
            
        # break
# rootdir='./种植问答网word'       
# list = os.listdir(rootdir) 
# for i in range(0,len(list)):
#     # print(i)
#     if i%1000==0:
#         print(i)
#     html_path = os.path.join(rootdir,list[i])
#     if os.path.isfile(html_path):
#         os.remove(html_path)
