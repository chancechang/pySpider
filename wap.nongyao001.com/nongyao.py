import xlrd
from bs4 import BeautifulSoup
import requests
import time
import datetime
import json
import re
import xlsxwriter
import os
from docx import Document
from docx.shared import Inches


cook='UM_distinctid=16822772df50-0cca678828aea6-b353461-100200-16822772e8fd5; CNZZDATA5626149=cnzz_eid%3D569227167-1546763274-%26ntime%3D1546779612'
header={
    'Cookie':cook,
    'Host':'m.my478.com',
    'Referer':'http://m.my478.com/jishu/list/',
    # 'Upgrade-Insecure-Requests':'1',
    'User-Agent':'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/64.0.3282.186 Safari/537.36',
}
#种植技术
def get_nongyao_link_one():
    f=open('nongyao_link.txt','w+',encoding='utf-8')
    url='https://wap.nongyao001.com/information/insects/index.html'
    req=requests.get(url)
    req.encoding='utf-8'
    bsObj=BeautifulSoup(req.text,'html.parser')
    jshu=bsObj.find_all('div',class_='right_content j-content')
    print(len(jshu))
    i=1
    for div in jshu:
        ul_list=div.find_all('ul')
        for ul in ul_list:
            a=ul.find('a')
            f.write(str([a.attrs['href'],i])+'\n')
        i=i+1
    f.close()
# get_nongyao_link_one()

def get_Obj(url):
    for i in range(5):
        try:
            req=requests.get(url)
            bsObj=BeautifulSoup(req.text,'html.parser')
            div=bsObj.find('div',class_='list_content')
            li_list=div.find_all('a',recursive=False)
            return li_list
        except:
            pass
        if i == 4:
            print(url+'失效')
    return None

def get_url(type,f):
    for i in range(1,1000):
        print(str(i)+'页')
        # https://wap.nongyao001.com/information/insects/list.html?code=6325&page=2
        url='https://wap.nongyao001.com'+str(type)+'&page='+str(i)
        a_list=get_Obj(url)
        print(len(a_list))
        for a in a_list:            
            href=a.attrs['href']
            title=a.find('h3').get_text().strip()
            row=[href,title]
            f.write(str(row)+'\n')
        if len(a_list)<10:
            break



def veg_get_all_link():   
    f=open('nongyao_link_two.txt','w+',encoding='utf-8')
    for line in open('nongyao_link.txt','r',encoding='utf-8'):
        line=eval(line)
        get_url(line[0],f)
    f.close()
# veg_get_all_link()


def d_load(src,imgrootname):
    root='./'+imgrootname
    path='./'+imgrootname+'/'+src.split('/')[-1]
    try:
        if not os.path.exists(path):
            if 'http' in src:
                r = requests.get(src)
            else:
                r = requests.get('https://wap.nongyao001.com'+src)
            r.raise_for_status()
            # 使用with语句可以不用自己手动关闭已经打开的文件流
            with open(path, "wb") as f:  # 开始写文件，wb代表写二进制文件
                f.write(r.content)
            print("爬取完成")
        else:
            # pass
            print("文件已存在")
    except Exception as e:
        print("爬取失败:"+str(e))
    
    return src.split('/')[-1]
def remove_control_characters(html):
    def str_to_int(s, default, base=10):
        if int(s, base) < 0x10000:
            return unichr(int(s, base))
        return default
        
    html = re.sub(u'&#(\d+);?'  ,lambda c: str_to_int(c.group(1), c.group(0)), html)

    html = re.sub(u"&#[xX]([0-9a-fA-F]+);?", lambda c: str_to_int(c.group(1), c.group(0), base=16), html)
    html = re.sub(u"[\x00-\x08\x0b\x0e-\x1f\x7f]", "", html)
    return html


def get_detail():
    f_error=open('nongyao_error.txt','w+',encoding='utf-8')
    f_img=open('nongyao_img.txt','w+',encoding='utf-8')
    # f_final=open('nongyao_final.txt','a+',encoding='utf-8')
    root='https://wap.nongyao001.com'
    wordrootname='农药网——病虫害word'
    htmlrootname='农药网——病虫害html'
    imgrootname='农药网——病虫害img'
    if not os.path.exists(wordrootname):
        os.mkdir(wordrootname)
    if not os.path.exists(htmlrootname):
        os.mkdir(htmlrootname)
    if not os.path.exists(imgrootname):
        os.mkdir(imgrootname)
    for line in open('nongyao_final.txt','r',encoding='utf-8'):
        line=eval(line)
        url=root+line[0] 
        # url='http://www.nongyao.com/newshow.asp?id=50163'
        # line[1]='白菜要丰产施肥必须有技法'
        print(url)   
        doc = Document() 

        content=BeautifulSoup(line[2],'html.parser') 
        #出去h3 和subStitute
        h3=content.find('h3')
        Subtitle=content.find('div',class_='Subtitle')
        content=BeautifulSoup(str(content).replace(str(h3),'').replace(str(Subtitle),''),'html.parser')
        img_list=content.find_all('img') 
        
        #content已经替换好标题和时间
        # 查找直接div和p
        div=content.find_all('div',recursive=False)
        div_p=content.find('p').find_all('div',recursive=False)
        p=content.find('p').find_all('p',recursive=False)
    
        if len(div)>=len(p) and len(div)>=len(div_p):
            repl=div
        elif len(p)>=len(div) and len(p)>=len(div_p):
            repl=p
        else:
            repl=div_p

        replace_list=[]
        print(len(repl))
        if len(repl)>5:
            for m in range(len(repl)-1,len(repl)-3,-1):
                if '咨询' in repl[m].get_text() or '了解更多' in repl[m].get_text() or '以上' in repl[m].get_text() or '参考' in repl[m].get_text() or '查看' in repl[m].get_text() or '还阅读了' in repl[m].get_text():
                    replace_list.append(str(repl[m]))
        if len(repl)!=0:
            tx=''.join('%s' %id for id in repl)
            for n in range(len(replace_list)):
                tx=tx.replace(replace_list[n],'') 
        else:
            tx=content
        #把tx变回bsObj
        tx_html=tx
        tx=BeautifulSoup(tx,'html.parser')          
           
        fname=line[1].replace('?','').replace('|','').replace('"','').replace('>','').replace('<','').replace('*','').replace('*','').replace('\\','').replace(':','').replace('/','')       
        path=fname+'_'+line[0][14:].replace('.html','') 

        doc.add_heading(line[1],level=0)
        for im in img_list:
            #写入图片        
            src=im.attrs['src']
            print(src)
            newSrc=d_load(src,imgrootname)
            tx_html=tx_html.replace(src,'../'+imgrootname+'/'+newSrc)
            try:
                doc.add_picture('./'+imgrootname+'/'+newSrc, width=Inches(3)) 
            except:
                pass
        
        #写入文本内容

        content_word=remove_control_characters(tx.get_text())
        doc.add_paragraph(u'%s'%(content_word)) 
        doc.save('./'+wordrootname+'/'+path+'.docx') 
    #写入html    
        
        fp=open('./'+htmlrootname+'/'+path+'.html','w',encoding='utf-8')
        fp.write('<html><head><meta http-equiv="Content-Type" content="text/html; charset=utf-8"></head><body><h1>'+str(line[1])+'</h1>'+str(tx_html)+'</body></html>') #写入数据           
        fp.close()
    f_error.close()
    f_img.close()


get_detail()


