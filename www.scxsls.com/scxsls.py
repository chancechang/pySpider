
from bs4 import BeautifulSoup
import requests
import time
import datetime
import json
import re
import xlsxwriter
import os
from docx import Document
from docx.shared import Inches


cook='UM_distinctid=16822772df50-0cca678828aea6-b353461-100200-16822772e8fd5; CNZZDATA5626149=cnzz_eid%3D569227167-1546763274-%26ntime%3D1546779612'
header={
    'Cookie':cook,
    'Host':'m.my478.com',
    'Referer':'http://m.my478.com/jishu/list/',
    # 'Upgrade-Insecure-Requests':'1',
    'User-Agent':'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/64.0.3282.186 Safari/537.36',
}
#种植技术
def get_scxsls_link():
    
    f=open('scxsls_link.txt','w+',encoding='utf-8')
    
    for i in range(1,649):
        print('第'+str(i)+'页')
        if i ==1:
            url='http://www.scxsls.com/anli/index.html'
        else:
            url='http://www.scxsls.com/anli/index_'+str(i)+'.html'
        req=requests.get(url)
        req.encoding='GB2312'
        bsObj=BeautifulSoup(req.text,'html.parser')
        jshu=bsObj.find_all('div',class_='cld_list')
        print(len(jshu))
        for div in jshu:
            a=div.find('h3').find('a')
            f.write(str([a.attrs['href'],a.get_text()])+'\n')
    f.close()


def get_Obj(url):
    for i in range(5):
        try:
            req=requests.get(url,headers=header)
            bsObj=BeautifulSoup(req.text,'html.parser')
            div=bsObj.find('ul',class_='bot_list')
            li_list=div.find_all('li')
            return li_list
        except:
            pass
        if i == 4:
            print(url+'失效')
    return None

def get_url(type,f):
    for i in range(1,1000):
        print(str(i)+'页')
        url='http://www.vegnet.com.cn'+str(type)+'_p'+str(i)+'.html'
        p_list=get_Obj(url)
        print(len(p_list))
        for p in p_list:
            a=p.find('a')
            href=a.attrs['href']
            title=a.attrs['title']
            row=[href,title]
            f.write(str(row)+'\n')
        if len(p_list)<12:
            break



def veg_get_all_link():   
    f=open('scxsls_line_two.txt','a+',encoding='utf-8')
    for line in open('scxsls_link.txt','r',encoding='utf-8'):
        line=eval(line)

        get_url(line[0].replace('.html',''),f)
    f.close()



def d_load(src,imgrootname):
    root='./'+imgrootname
    path='./'+imgrootname+'/'+src.split('/')[-1]
    try:
        if not os.path.exists(path):
            r = requests.get(src)
            r.raise_for_status()
            # 使用with语句可以不用自己手动关闭已经打开的文件流
            with open(path, "wb") as f:  # 开始写文件，wb代表写二进制文件
                f.write(r.content)
            print("爬取完成")
        else:
            print("文件已存在")
    except Exception as e:
        print("爬取失败:"+str(e))
    
    return src.split('/')[-1]
def remove_control_characters(html):
    def str_to_int(s, default, base=10):
        if int(s, base) < 0x10000:
            return unichr(int(s, base))
        return default
        
    html = re.sub(u'&#(\d+);?'  ,lambda c: str_to_int(c.group(1), c.group(0)), html)

    html = re.sub(u"&#[xX]([0-9a-fA-F]+);?", lambda c: str_to_int(c.group(1), c.group(0), base=16), html)
    html = re.sub(u"[\x00-\x08\x0b\x0e-\x1f\x7f]", "", html)
    return html


def get_detail():
    f_error=open('scxsls_error.txt','a+',encoding='utf-8')
    f_final=open('scxsls_final.txt','a+',encoding='utf-8')
    wordrootname='刑事案例word'
    if not os.path.exists(wordrootname):
        os.mkdir(wordrootname)
    r=0
    for line in open('scxsls_link.txt','r',encoding='utf-8'):
        line=eval(line)
        if 'http' in line[0]:
            url=line[0] 
        else:
            url='http://www.scxsls.com'+line[0]
        id=url.split('/')[-1].replace('.html','')
        print(url)
        fname=line[1].strip().replace('?','').replace('|','').replace('"','').replace('>','').replace('<','').replace('*','').replace('*','').replace('\\','').replace(':','').replace('/','').replace('\t','').replace('\r','').replace('\n','')
           
        path1='./'+wordrootname+'/'+fname+'_'+str(id)
        # print(path1)
        
        if os.path.exists(path1+'.docx'):
            print('已存在')
            continue
        
        
        # print(id)
        doc = Document()  
        for i in range(3):
            try:
                req=requests.get(url,timeout=15)
                req.encoding='GB2312'
                bsObj=BeautifulSoup(req.text,'lxml')
                #
                content=bsObj.find('div',{'id':'news_content'})
                origin_text=bsObj.find('div',{'id':'news_meta_left'}).get_text()
                #写入标题
                break
            except:
                pass

        # print(origin_text)
        # print(content)
        # break
        if content==None or '来源' not in origin_text :
            print(url+'未取到数据')
            f_error.write(str(line)+'\n')
            continue
       
        content_word=str(content).replace('</p','</p')
        
        print(fname)    
        doc.add_heading(line[1],level=0)
        origin_text=origin_text.split('\u3000')
        for ori in origin_text:
            if '来源' in ori:
                origin_text=ori
                break
        try:
            origin_text=re.search(r'来源：(.+?)\| 作者',origin_text).group(1)
        except:
            pass
   
        if '来源' in origin_text:
            pass
        else:
            origin_text='来源：'+origin_text
        doc.add_paragraph(u'%s'%(origin_text)) 
        print(origin_text)
                
        # print(origin_text)        
        #写入文本内容
        # content_word=remove_control_characters(BeautifulSoup(content_word,'html.parser').get_text())   
        content_word=BeautifulSoup(content_word,'html.parser').get_text()  
        doc.add_paragraph(u'%s'%( content_word)) 
        doc.save(path1+'.docx') 
        # r=r+1
        # if r==5:
        #     break
        # break
    f_error.close()
    f_final.close()

get_detail()

