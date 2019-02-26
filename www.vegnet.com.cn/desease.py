
from bs4 import BeautifulSoup
import requests
import time
import datetime
import json
import re
import xlsxwriter
import os
from docx import Document
from docx.shared import Inches


cook='yjs_ab_lid=8f40f131b0d3a4a7a735b95021dccb7656430; yjs_ab_score=300; __cfduid=d109b4638e427a611aef1ad3aa8c45a931546651280; yjs_id=db55d5dba4e5e03a59dee0b9239ae3a6; ctrl_time=1; bdshare_firstime=1546651295232; Hm_lvt_eda63048a2a5f2f971045c7b18e63a72=1546651295; __utmc=247588352; __utmz=247588352.1546651297.1.1.utmcsr=(direct)|utmccn=(direct)|utmcmd=(none); Hm_lvt_8f396c620ee16c3d331cf84fd54feafe=1546679303; cf_clearance=c73d79a565e1bfd12d975ad71ff97d73cd4d47ab-1546737527-31536000-150; Hm_lpvt_8f396c620ee16c3d331cf84fd54feafe=1546737529; __utma=247588352.1427397888.1546651297.1546700969.1546737584.7; __utmt=1; Hm_lpvt_eda63048a2a5f2f971045c7b18e63a72=1546737602; __utmb=247588352.3.10.1546737584'
header={
    'Cookie':cook,
    'Host':'www.vegnet.com.cn',
    'Referer':'http://www.vegnet.com.cn/Tech/List_oc3135.html',
    'Upgrade-Insecure-Requests':'1',
    'User-Agent':'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/64.0.3282.186 Safari/537.36',
}
def get_link_one():
    f=open('desease_line_one.txt','a+',encoding='utf-8')
    url='http://www.vegnet.com.cn/Disease/List_c1025.html'
    req=requests.get(url,headers=header)
    bsObj=BeautifulSoup(req.text,'html.parser')
    a_list=bsObj.find('div',class_='lb2').find_all('a')
    print(len(a_list))
    for a in a_list:
        f.write(str([a.attrs['href'],a.get_text()])+'\n')
    f.close()

def get_Obj(url):
    for i in range(5):
        try:
            req=requests.get(url,headers=header)
            bsObj=BeautifulSoup(req.text,'html.parser')
            div=bsObj.find('ul',class_='bot_list')
            li_list=div.find_all('li')
            return li_list
        except:
            pass
        if i == 4:
            print(url+'失效')
    return None

def get_url(type,f):
    for i in range(1,1000):
        print(str(i)+'页')
        url='http://www.vegnet.com.cn'+str(type)+'_p'+str(i)+'.html'
        p_list=get_Obj(url)
        print(len(p_list))
        for p in p_list:
            a=p.find('a')
            href=a.attrs['href']
            title=a.attrs['title']
            row=[href,title]
            f.write(str(row)+'\n')
        if len(p_list)<12:
            break



def veg_get_all_link():   
    f=open('desease_line_two.txt','a+',encoding='utf-8')
    for line in open('desease_line_one.txt','r',encoding='utf-8'):
        line=eval(line)

        get_url(line[0].replace('.html',''),f)
    f.close()



def d_load(src,imgrootname):
    root='./'+imgrootname
    path='./'+imgrootname+'/'+src.split('/')[-1]
    try:
        if not os.path.exists(path):
            r = requests.get(src)
            r.raise_for_status()
            # 使用with语句可以不用自己手动关闭已经打开的文件流
            with open(path, "wb") as f:  # 开始写文件，wb代表写二进制文件
                f.write(r.content)
            print("爬取完成")
        else:
            print("文件已存在")
    except Exception as e:
        print("爬取失败:"+str(e))
    
    return src.split('/')[-1]
def remove_control_characters(html):
    def str_to_int(s, default, base=10):
        if int(s, base) < 0x10000:
            return unichr(int(s, base))
        return default
        
    html = re.sub(u'&#(\d+);?'  ,lambda c: str_to_int(c.group(1), c.group(0)), html)

    html = re.sub(u"&#[xX]([0-9a-fA-F]+);?", lambda c: str_to_int(c.group(1), c.group(0), base=16), html)
    html = re.sub(u"[\x00-\x08\x0b\x0e-\x1f\x7f]", "", html)
    return html


def get_detail():
 
    f_error=open('desease_error.txt','w+',encoding='utf-8')
    f_img=open('desease_img.txt','w+',encoding='utf-8')
    root='http://www.vegnet.com.cn'
    wordrootname='瓜类病虫害word'
    htmlrootname='瓜类病虫害html'
    imgrootname='瓜类病虫害img'
    if not os.path.exists(wordrootname):
        os.mkdir(wordrootname)
    if not os.path.exists(htmlrootname):
        os.mkdir(htmlrootname)
    if not os.path.exists(imgrootname):
        os.mkdir(imgrootname)
    for line in open('desease_line_two.txt','r',encoding='utf-8'):
        line=eval(line)
        url=root+line[0] 
        print(url)   
        doc = Document()  
        for i in range(3):
            try:
                req=requests.get(url,headers=header,timeout=15)
                req.encoding='utf-8'
                bsObj=BeautifulSoup(req.text,'html.parser')
                #病虫害
                content=bsObj.find('div',class_='detai_main')
                skin_green=bsObj.find('div',class_='OriginalPicBorder')
                skin_text=str(skin_green)
                img_list=skin_green.find_all('img')
                #写入标题
                break
            except:
                pass
        # print(content)
        if content==None  :
            print(url+'未取到数据')
            f_error.write(str(line)+'\n')
            continue
        fname=line[1].strip().replace('?','').replace('|','').replace('"','').replace('>','').replace('<','').replace('*','').replace('*','').replace('\\','').replace(':','').replace('/','')
        
        path='./'+wordrootname+'/'+fname+'.docx'
        
        print(fname)    
      
                    
        doc.add_heading(line[1],level=0)
        #写入图片
        
        for im in img_list:
            src=im.attrs['src']
            print(src)
            newSrc=d_load(src,imgrootname)
            #替换内容
            skin_text=skin_text.replace(src,'../'+imgrootname+'/'+newSrc)
            try:
                doc.add_picture('./'+imgrootname+'/'+newSrc, width=Inches(3)) 
            except:
                pass
            
        #写入文本内容
        # content_word=remove_control_characters(str(content.get_text()))   
        doc.add_paragraph(u'%s'%(str(content.get_text()))) 
        doc.save('./'+wordrootname+'/'+fname+'.docx') 
    #写入html    
        
        fp=open('./'+htmlrootname+'/'+fname+'.html','w',encoding='utf-8')
        fp.write('<html><head><meta http-equiv="Content-Type" content="text/html; charset=utf-8"></head><body><h1>'+str(line[1])+'</h1>'+str(skin_text)+str(content)+'</body></html>') #写入数据
        fp.close()
        # break
    f_error.close()
    f_img.close()

get_detail()


# rootdir = './蔬菜技术html'
# list = os.listdir(rootdir) #列出文件夹下所有的目录与文件
# for i in range(0,len(list)):
#     print(i)
#     path = os.path.join(rootdir,list[i])
#     if os.path.isfile(path):
#         tex=open(path,'r',encoding='utf-8')
#         if 'img' in tex.read():
#             print(path)
#             break
       


