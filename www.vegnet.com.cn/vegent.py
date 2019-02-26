
from bs4 import BeautifulSoup
import requests
import time
import datetime
import json
import re
import xlsxwriter
import os
from docx import Document
from docx.shared import Inches


cook='yjs_ab_lid=8f40f131b0d3a4a7a735b95021dccb7656430; yjs_ab_score=300; __cfduid=d109b4638e427a611aef1ad3aa8c45a931546651280; yjs_id=db55d5dba4e5e03a59dee0b9239ae3a6; ctrl_time=1; bdshare_firstime=1546651295232; Hm_lvt_eda63048a2a5f2f971045c7b18e63a72=1546651295; __utmc=247588352; __utmz=247588352.1546651297.1.1.utmcsr=(direct)|utmccn=(direct)|utmcmd=(none); __utma=247588352.1427397888.1546651297.1546657857.1546674922.4; cf_clearance=85fd175670264be5dbd7746c9fa4a2c990df25ee-1546676877-31536000-150; Hm_lpvt_eda63048a2a5f2f971045c7b18e63a72=1546676878; __utmt=1; __utmb=247588352.5.10.1546674922'
cook='yjs_ab_lid=8f40f131b0d3a4a7a735b95021dccb7656430; yjs_ab_score=300; __cfduid=d109b4638e427a611aef1ad3aa8c45a931546651280; yjs_id=db55d5dba4e5e03a59dee0b9239ae3a6; ctrl_time=1; bdshare_firstime=1546651295232; Hm_lvt_eda63048a2a5f2f971045c7b18e63a72=1546651295; __utmc=247588352; __utmz=247588352.1546651297.1.1.utmcsr=(direct)|utmccn=(direct)|utmcmd=(none); Hm_lvt_8f396c620ee16c3d331cf84fd54feafe=1546679303; Hm_lpvt_8f396c620ee16c3d331cf84fd54feafe=1546679303; cf_clearance=a75a9fadab5e4bd9764f32f3bc60c1dd2eba0f94-1546681733-31536000-150; Hm_lpvt_eda63048a2a5f2f971045c7b18e63a72=1546681734; __utma=247588352.1427397888.1546651297.1546674922.1546681735.5; __utmt=1; __utmb=247588352.1.10.1546681735'
cook='yjs_ab_lid=8f40f131b0d3a4a7a735b95021dccb7656430; yjs_ab_score=300; __cfduid=d109b4638e427a611aef1ad3aa8c45a931546651280; yjs_id=db55d5dba4e5e03a59dee0b9239ae3a6; ctrl_time=1; bdshare_firstime=1546651295232; Hm_lvt_eda63048a2a5f2f971045c7b18e63a72=1546651295; __utmc=247588352; __utmz=247588352.1546651297.1.1.utmcsr=(direct)|utmccn=(direct)|utmcmd=(none); Hm_lvt_8f396c620ee16c3d331cf84fd54feafe=1546679303; cf_clearance=c73d79a565e1bfd12d975ad71ff97d73cd4d47ab-1546737527-31536000-150; Hm_lpvt_8f396c620ee16c3d331cf84fd54feafe=1546737529; __utma=247588352.1427397888.1546651297.1546700969.1546737584.7; __utmt=1; Hm_lpvt_eda63048a2a5f2f971045c7b18e63a72=1546737602; __utmb=247588352.3.10.1546737584'
header={
    'Cookie':cook,
    'Host':'www.vegnet.com.cn',
    'Referer':'http://www.vegnet.com.cn/Tech/List_oc3135.html',
    'Upgrade-Insecure-Requests':'1',
    'User-Agent':'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/64.0.3282.186 Safari/537.36',
}

def get_Obj(url):
    for i in range(5):
        try:
            req=requests.get(url,headers=header)
            bsObj=BeautifulSoup(req.text,'html.parser')
            div=bsObj.find('div',class_='jxs_list jxs_list1')
            p_list=div.find_all('p')
            return p_list
        except:
            pass
        if i == 4:
            print(url+'失效')
    return None

def get_url(type,f):
    for i in range(1,1000):
        print(str(i)+'页')
        url='http://www.vegnet.com.cn/Tech/List_oc'+str(type)+'_p'+str(i)+'.html'
        p_list=get_Obj(url)
        print(len(p_list))
        for p in p_list:
            href=p.find('a').attrs['href']
            orgin=p.find('span',class_='k_159')
            row=[href,str(orgin),p.find('a').get_text()]
            f.write(str(row)+'\n')
        if len(p_list)<20:
            break
def veg_get_all_link():   
    f=open('cabbage_final.txt','a+',encoding='utf-8')
    for type in range(3135,3166):
        print('种类'+str(type))
        get_url(type,f)
    f.close()



def d_load(src,imgrootname):
    root='./'+imgrootname
    path='./'+imgrootname+'/'+src.split('/')[-1]
    try:
        if not os.path.exists(path):
            r = requests.get(src)
            r.raise_for_status()
            # 使用with语句可以不用自己手动关闭已经打开的文件流
            with open(path, "wb") as f:  # 开始写文件，wb代表写二进制文件
                f.write(r.content)
            print("爬取完成")
        else:
            print("文件已存在")
    except Exception as e:
        print("爬取失败:"+str(e))
    
    return src.split('/')[-1]
def remove_control_characters(html):
    def str_to_int(s, default, base=10):
        if int(s, base) < 0x10000:
            return unichr(int(s, base))
        return default
        
    html = re.sub(u'&#(\d+);?'  ,lambda c: str_to_int(c.group(1), c.group(0)), html)

    html = re.sub(u"&#[xX]([0-9a-fA-F]+);?", lambda c: str_to_int(c.group(1), c.group(0), base=16), html)
    html = re.sub(u"[\x00-\x08\x0b\x0e-\x1f\x7f]", "", html)
    return html


def get_detail():
 
    f_error=open('cabbage_error2.txt','w+',encoding='utf-8')
    f_img=open('cabbage_img.txt','w+',encoding='utf-8')
    root='http://www.vegnet.com.cn'
    wordrootname='蔬菜技术word'
    htmlrootname='蔬菜技术html'
    imgrootname='蔬菜技术img'
    if not os.path.exists(wordrootname):
        os.mkdir(wordrootname)
    if not os.path.exists(htmlrootname):
        os.mkdir(htmlrootname)
    if not os.path.exists(imgrootname):
        os.mkdir(imgrootname)
    for line in open('cabbage_img1.txt','r',encoding='utf-8'):
        line=eval(line)
        line1_html='来源：'+line[1]
        line1_word='来源：'+str(BeautifulSoup(line[1],'html.parser').get_text()).strip().replace('\n','')
        
        # fname=line[2].strip().replace('?','').replace('|','').replace('"','').replace('>','').replace('<','').replace('*','').replace('*','').replace('\\','').replace(':','').replace('/','')
        
        # path='./'+wordrootname+'/'+fname+'.docx'
        
        # if  os.path.exists(path):
        #     continue
        # else:
        #     f_img.write(str(line)+'\n')
        #     continue
        
        if 'vegnet'  in line[0]:
            url=line[0]
        else:
            url=root+line[0]
        # url='http://www.vegnet.com.cn/Disease/2862.html?cateID=1025'
        print(url)
        
        doc = Document()  
        for i in range(1):
            try:
                req=requests.get(url,headers=header,timeout=15)
                req.encoding='utf-8'
                bsObj=BeautifulSoup(req.text,'html.parser')
                h1=bsObj.find('h1',{'id':'newstitle'})
                #蔬菜
                content=bsObj.find('div',class_='main_news')
                img_list=content.find_all('img')
                # print(bsObj)

                
                #病虫害
                # content=bsObj.find('div',class_='detai_main')
                # skin_green=bsObj.find('div',class_='OriginalPicBorder')
                # skin_text=str(skin_green)
                # img_list=skin_green.find_all('img')
                #写入标题
                break
            except:
                
                try:
                    h1=bsObj.find('div',{'id':'newstitle'}).find('h1')
                    content=bsObj.find('div',{'id':'fontzoom'})
                    img_list=content.find_all('img')
              
                except:
                    try:
                        h1=bsObj.find('h3')
                        content=bsObj.find('div',{'id':'fontzoom'})
                        img_list=content.find_all('img')
                    except:
                        pass
        # print(content)
        if content==None or h1==None :
            print(url+'未取到数据')
            f_error.write(str(line)+'\n')
            continue
        # if len(img_list)>0:
        #     print(url+'存在图片')
        #     f_img.write(str(line)+'\n')
        #     continue
        fname=h1.get_text().strip().replace('?','').replace('|','').replace('"','').replace('>','').replace('<','').replace('*','').replace('*','').replace('\\','').replace(':','').replace('/','')
        
        path='./'+wordrootname+'/'+fname+'.docx'
        
        print(fname)    
        print('写入')    
                    
        doc.add_heading(h1.get_text(),level=0)
        doc.add_paragraph(line1_word)
        #写入图片
        content_html=str(content)
        for im in img_list:
            src=im.attrs['src']
            print(src)
            newSrc=d_load(src,imgrootname)
            #替换内容
            # skin_text=skin_text.replace(src,newSrc)
          
            content_html=str(content).replace(src,'../'+imgrootname+'/'+newSrc)
            try:
                doc.add_picture('./'+imgrootname+'/'+newSrc, width=Inches(3)) 
            except:
                pass 
        #写入文本内容
        content_word=remove_control_characters(str(content.get_text()))   
        doc.add_paragraph(u'%s'%(content_word)) 
        doc.save('./'+wordrootname+'/'+fname+'.docx') 
    #写入html    
        
        fp=open('./'+htmlrootname+'/'+fname+'.html','w',encoding='utf-8')
        fp.write('<html><head><meta http-equiv="Content-Type" content="text/html; charset=utf-8"></head><body>'+str(h1)+str(line1_html )+str(content_html)+'</body></html>') #写入数据
        fp.close()
        # break
    f_error.close()
    f_img.close()


get_detail()

# rootdir = './蔬菜技术html'
# list = os.listdir(rootdir) #列出文件夹下所有的目录与文件
# for i in range(0,len(list)):
#     print(i)
#     path = os.path.join(rootdir,list[i])
#     if os.path.isfile(path):
#         tex=open(path,'r',encoding='utf-8')
#         if 'img' in tex.read():
#             print(path)
#             break
       


