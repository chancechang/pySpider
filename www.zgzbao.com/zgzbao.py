import xlrd
from bs4 import BeautifulSoup
import requests
import time
import datetime
import json
import re
import xlsxwriter
import os
from docx import Document
from docx.shared import Inches


cook='UM_distinctid=16822772df50-0cca678828aea6-b353461-100200-16822772e8fd5; CNZZDATA5626149=cnzz_eid%3D569227167-1546763274-%26ntime%3D1546779612'
header={
    'Cookie':cook,
    'Host':'m.my478.com',
    'Referer':'http://m.my478.com/jishu/list/',
    # 'Upgrade-Insecure-Requests':'1',
    'User-Agent':'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/64.0.3282.186 Safari/537.36',
}
#种植技术
def get_zgzbao_link_one():
    f=open('zgzbao_link.txt','w+',encoding='utf-8')
    for i in range(1,146):
        url='http://www.zgzbao.com/jscs.asp?page='+str(i)+'&bigclass=%BC%BC%CA%F5%B3%AC%CA%D0&smallclass='
        
        req=requests.get(url)
        req.encoding='gb2312'
        bsObj=BeautifulSoup(req.text,'html.parser')
        jshu=bsObj.find_all('td',{'width':874})
        print(len(jshu))
        for td in jshu:
            a=td.find('a')
            f.write(str([a.attrs['href'],a.get_text()])+'\n')
    f.close()


def get_Obj(url):
    for i in range(5):
        try:
            req=requests.get(url,headers=header)
            bsObj=BeautifulSoup(req.text,'html.parser')
            div=bsObj.find('ul',class_='bot_list')
            li_list=div.find_all('li')
            return li_list
        except:
            pass
        if i == 4:
            print(url+'失效')
    return None

def get_url(type,f):
    for i in range(1,1000):
        print(str(i)+'页')
        url='http://www.vegnet.com.cn'+str(type)+'_p'+str(i)+'.html'
        p_list=get_Obj(url)
        print(len(p_list))
        for p in p_list:
            a=p.find('a')
            href=a.attrs['href']
            title=a.attrs['title']
            row=[href,title]
            f.write(str(row)+'\n')
        if len(p_list)<12:
            break



def veg_get_all_link():   
    f=open('zgzbao_line_two.txt','a+',encoding='utf-8')
    for line in open('zgzbao_line_one.txt','r',encoding='utf-8'):
        line=eval(line)

        get_url(line[0].replace('.html',''),f)
    f.close()



def d_load(src,imgrootname):
    root='./'+imgrootname
    path='./'+imgrootname+'/'+src.split('/')[-1]
    try:
        if not os.path.exists(path):
            if 'http' in src:
                r = requests.get(src)
            else:
                r = requests.get('http://www.zgzbao.com/'+src)
            r.raise_for_status()
            # 使用with语句可以不用自己手动关闭已经打开的文件流
            with open(path, "wb") as f:  # 开始写文件，wb代表写二进制文件
                f.write(r.content)
            print("爬取完成")
        else:
            print("文件已存在")
    except Exception as e:
        print("爬取失败:"+str(e))
    
    return src.split('/')[-1]
def remove_control_characters(html):
    def str_to_int(s, default, base=10):
        if int(s, base) < 0x10000:
            return unichr(int(s, base))
        return default
        
    html = re.sub(u'&#(\d+);?'  ,lambda c: str_to_int(c.group(1), c.group(0)), html)

    html = re.sub(u"&#[xX]([0-9a-fA-F]+);?", lambda c: str_to_int(c.group(1), c.group(0), base=16), html)
    html = re.sub(u"[\x00-\x08\x0b\x0e-\x1f\x7f]", "", html)
    return html


def get_detail():
    f_error=open('zgzbao_error1.txt','w+',encoding='utf-8')
    f_img=open('zgzbao_img1.txt','w+',encoding='utf-8')
    f_final=open('zgzbao_final.txt','a+',encoding='utf-8')
    root='http://www.zgzbao.com/'
    wordrootname='植保网_技术超市word'
    htmlrootname='植保网_技术超市html'
    imgrootname='植保网_技术超市img'
    if not os.path.exists(wordrootname):
        os.mkdir(wordrootname)
    if not os.path.exists(htmlrootname):
        os.mkdir(htmlrootname)
    if not os.path.exists(imgrootname):
        os.mkdir(imgrootname)
    for line in open('repeat.txt','r',encoding='utf-8'):
        line=eval(line)
        url=root+line[0] 
        # url='http://www.zgzbao.com/newshow.asp?id=50163'
        # line[1]='白菜要丰产施肥必须有技法'
        print(url)   
        doc = Document()  
        content=BeautifulSoup(line[2],'html.parser')
        img_list=content.find_all('img')
        # for i in range(3):
        #     try:
        #         req=requests.get(url,headers=header,timeout=15)
        #         req.encoding='gb18030'
        #         bsObj=BeautifulSoup(req.text,'html.parser')
        #         #病虫害
        #         content=bsObj.find('div',{'id':'Zoom'})
        #         img_list=content.find_all('img')
        #         #写入标题
        #         break
        #     except:
        #         pass
        if content==None  :
            print(url+'未取到数据')
            f_error.write(str(line)+'\n')
            continue
        # if len(img_list)>0:
        #     print(url+'存在图片')
        #     f_img.write(str(line)+'\n')
        #     continue
        line[1]=line[1].replace('\r\n','').strip()
        # f_final.write(str(line+[str(content)])+'\n')
        # continue        
        
        fname=line[1].replace('?','').replace('|','').replace('"','').replace('>','').replace('<','').replace('*','').replace('*','').replace('\\','').replace(':','').replace('/','')       
        # path='./'+wordrootname+'/'+fname+'_'+line[0][15:]+'.docx'
        # print(path)
        
        # print(fname)    
        # doc.add_heading(line[1],level=0)
        # #写入图片
        # print(content.get_text())
        # skin_text=str(content)
        # for im in img_list:
        #     src=im.attrs['src']
        #     print(src)
        #     newSrc=d_load(src,imgrootname)
        #     #替换内容
        #     skin_text=skin_text.replace(src,'../'+imgrootname+'/'+newSrc)
        #     try:
        #         doc.add_picture('./'+imgrootname+'/'+newSrc, width=Inches(3)) 
        #     except:
        #         pass
            
        #写入文本内容
        # content_word=remove_control_characters(str(content.get_text()))   
        # doc.add_paragraph(u'%s'%( content_word)) 
        # doc.save(path) 
    #写入html    
        
        # fp=open('./'+htmlrootname+'/'+fname+'_'+line[0][15:]+'.html','w',encoding='gb18030')
        try:
            os.remove('./'+htmlrootname+'/'+fname+'.html')
            
        except:
            pass
        try:
            os.remove('./'+wordrootname+'/'+fname+'.docx')
        except:
            pass
      
        # fp.write('<html><head><meta http-equiv="Content-Type" content="text/html; charset=gb2312"></head><body><h1>'+str(line[1])+'</h1>'+str(skin_text)+'</body></html>') #写入数据
            
        # fp.close()
        # break
    f_error.close()
    f_img.close()
    f_final.close()

# get_detail()



wordrootname='植保网_技术超市word'
htmlrootname='植保网_技术超市html'
imgrootname='植保网_技术超市img'
rootdir = './植保网_技术超市html'
if not os.path.exists(wordrootname):
    os.mkdir(wordrootname)
list = os.listdir(rootdir) #列出文件夹下所有的目录与文件
for i in range(0,len(list)):
    print(i)
    
    print(list[i])
    html_path = os.path.join(rootdir,list[i])
    fname=list[i].replace('.html','')
    if os.path.isfile(html_path):
        doc = Document()
        doc.add_heading(fname,level=0)
        tex=open(html_path,'r',encoding='gb18030')
        bsObj=BeautifulSoup(tex.read(),'html.parser')
        img_list=bsObj.find_all('img')
        content=str(bsObj).replace('</p','\n</p')
        for  im in img_list:
            try:
                print(im.attrs['src'][1:])
                
                doc.add_picture(im.attrs['src'][1:], ) 
            except:
                print(os.path.isfile(im.attrs['src'][1:]))
        content_word=remove_control_characters(BeautifulSoup(content,'html.parser').get_text())
        doc.add_paragraph(u'%s'%(content_word.replace(fname,''))) 
        doc.save('./'+wordrootname+'/'+fname+'.docx') 
    # break


       


